from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx

doc = Document()

# Title
title = doc.add_heading('End-to-End CRUD Operations Demonstration', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

sections = [
    {
        "title": "1. Create: New Record Added",
        "caption": "Caption: The user navigated to the frontend interface and successfully created a new Task (e.g., 'Complete Mathematics Assignment'). The frontend sent the data to the Supabase database, and the new record instantly appeared in the active task list.",
        "placeholder": "[ PASTE YOUR SCREENSHOT OF THE NEWLY CREATED TASK HERE ]"
    },
    {
        "title": "2. Read: List of Records Displayed",
        "caption": "Caption: The user navigated to the Dashboard/Today module. The frontend queried the Supabase database and successfully retrieved the list of all active tasks, displaying them clearly in the UI task list.",
        "placeholder": "[ PASTE YOUR SCREENSHOT OF THE FULL TASK LIST HERE ]"
    },
    {
        "title": "3. Update: Record Updated Successfully",
        "caption": "Caption: The user interacted with a specific task by clicking the completion checkbox (or editing its details). The frontend successfully updated the record's state in the database, and the UI immediately updated to reflect the completed/modified status.",
        "placeholder": "[ PASTE YOUR SCREENSHOT OF THE UPDATED/COMPLETED TASK HERE ]"
    },
    {
        "title": "4. Delete: Record Removed",
        "caption": "Caption: The user selected an existing record and triggered the delete action. The frontend sent a delete request to the database, successfully removing the record and instantly clearing it from the user's view.",
        "placeholder": "[ PASTE YOUR SCREENSHOT OF THE LIST AFTER DELETION HERE ]"
    }
]

for section in sections:
    doc.add_heading(section["title"], level=1)
    
    p_caption = doc.add_paragraph()
    p_caption.add_run(section["caption"]).bold = False
    
    p_placeholder = doc.add_paragraph()
    p_placeholder.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p_placeholder.add_run(section["placeholder"])
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0) # Red color to make it obvious
    
    doc.add_page_break()

doc.save('CRUD_Operations_Demonstration.docx')
print("Document created successfully.")
